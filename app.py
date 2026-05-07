import os
import re
import json
import time
import textwrap
from datetime import datetime
from urllib.parse import urlparse

import requests
import pandas as pd
import streamlit as st
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from docx import Document
from pypdf import PdfReader
from rapidfuzz import fuzz

try:
    from openai import OpenAI
except Exception:
    OpenAI = None


# ============================================================
# CONFIG
# ============================================================

load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "") or st.secrets.get("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "") or st.secrets.get("OPENAI_MODEL", "gpt-4.1-mini")
SERPAPI_API_KEY = os.getenv("SERPAPI_API_KEY", "") or st.secrets.get("SERPAPI_API_KEY", "")

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ============================================================
# BASIC HELPERS
# ============================================================

def clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def safe_filename(text: str) -> str:
    text = re.sub(r"[^a-zA-Z0-9_\- ]", "", text)
    text = re.sub(r"\s+", "_", text)
    return text[:80]


def get_domain(url: str) -> str:
    try:
        parsed = urlparse(url)
        domain = parsed.netloc.lower()
        domain = domain.replace("www.", "")
        return domain
    except Exception:
        return ""


def truncate(text: str, limit: int = 5000) -> str:
    if not text:
        return ""
    return text[:limit]


def call_openai(prompt: str, max_retries: int = 2) -> str:
    """
    Simple OpenAI wrapper.
    If no API key is present, returns empty string so the app can still run partially.
    """
    if not OPENAI_API_KEY or OpenAI is None:
        return ""

    client = OpenAI(api_key=OPENAI_API_KEY)

    for attempt in range(max_retries + 1):
        try:
            response = client.responses.create(
                model=OPENAI_MODEL,
                input=prompt,
            )
            return response.output_text.strip()
        except Exception as e:
            if attempt == max_retries:
                return f"OPENAI_ERROR: {str(e)}"
            time.sleep(1.5)

    return ""


def extract_json_from_text(text: str):
    """
    Tries to extract JSON even if the model wraps it in text.
    """
    if not text:
        return None

    text = text.strip()

    try:
        return json.loads(text)
    except Exception:
        pass

    match = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except Exception:
            return None

    return None


# ============================================================
# CV PARSING
# ============================================================

def read_pdf_cv(uploaded_file) -> str:
    reader = PdfReader(uploaded_file)
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pass
    return clean_text("\n".join(pages))


def read_docx_cv(uploaded_file) -> str:
    doc = Document(uploaded_file)
    paragraphs = [p.text for p in doc.paragraphs]
    return clean_text("\n".join(paragraphs))


def read_txt_cv(uploaded_file) -> str:
    return uploaded_file.read().decode("utf-8", errors="ignore")


def read_uploaded_cv(uploaded_file) -> str:
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return read_pdf_cv(uploaded_file)

    if name.endswith(".docx"):
        return read_docx_cv(uploaded_file)

    if name.endswith(".txt"):
        return read_txt_cv(uploaded_file)

    raise ValueError("Unsupported CV format. Please upload PDF, DOCX, or TXT.")


def build_candidate_profile(cv_text: str) -> dict:
    prompt = f"""
You are an expert CV parser.

Extract the candidate profile from this CV.

Return ONLY valid JSON with this schema:

{{
  "name": "",
  "email": "",
  "phone": "",
  "location": "",
  "education": [],
  "experience": [],
  "projects": [],
  "technical_skills": [],
  "soft_skills": [],
  "certifications": [],
  "target_keywords": []
}}

Rules:
- Do not invent anything.
- Use only information explicitly present in the CV.
- Keep bullets concise.
- If something is missing, use an empty string or empty list.

CV:
\"\"\"
{truncate(cv_text, 12000)}
\"\"\"
"""

    result = call_openai(prompt)
    parsed = extract_json_from_text(result)

    if parsed:
        return parsed

    # Fallback if no OpenAI key or parsing fails.
    return {
        "name": "",
        "email": "",
        "phone": "",
        "location": "",
        "education": [],
        "experience": [],
        "projects": [],
        "technical_skills": simple_skill_extract(cv_text),
        "soft_skills": [],
        "certifications": [],
        "target_keywords": simple_skill_extract(cv_text),
    }


def simple_skill_extract(text: str) -> list:
    common_skills = [
        "python", "pandas", "numpy", "scikit-learn", "sklearn", "pytorch",
        "tensorflow", "machine learning", "deep learning", "nlp", "llm",
        "transformers", "bert", "finbert", "sql", "aws", "azure", "gcp",
        "docker", "linux", "cybersecurity", "soc", "siem", "splunk",
        "data analysis", "data science", "statistics", "matplotlib",
        "seaborn", "fastapi", "streamlit", "git", "github"
    ]

    found = []
    lower = text.lower()

    for skill in common_skills:
        if skill in lower:
            found.append(skill)

    return sorted(set(found))


# ============================================================
# JOB SEARCH
# ============================================================

def serpapi_search(query: str, num_results: int = 10) -> list:
    """
    Uses SerpAPI Google search.
    """
    if not SERPAPI_API_KEY:
        st.warning("SERPAPI_API_KEY is missing. Add it to your .env file.")
        return []

    url = "https://serpapi.com/search.json"

    params = {
        "engine": "google",
        "q": query,
        "api_key": SERPAPI_API_KEY,
        "num": num_results,
        "hl": "en",
        "gl": "uk",
    }

    try:
        response = requests.get(url, params=params, timeout=20)
        response.raise_for_status()
        data = response.json()
    except Exception as e:
        st.error(f"Search failed for query: {query}\nError: {e}")
        return []

    results = []

    for item in data.get("organic_results", []):
        results.append({
            "title": clean_text(item.get("title", "")),
            "url": item.get("link", ""),
            "snippet": clean_text(item.get("snippet", "")),
            "source": get_domain(item.get("link", "")),
            "query": query,
        })

    return results


def generate_job_queries(target_roles: str, location: str, extra_keywords: str) -> list:
    roles = [r.strip() for r in re.split(r",|\n", target_roles) if r.strip()]

    queries = []

    for role in roles:
        queries.extend([
            f'"{role}" "{location}" graduate jobs',
            f'"{role}" "{location}" internship',
            f'"{role}" "{location}" early careers',
            f'"{role}" "{location}" site:targetjobs.co.uk',
            f'"{role}" "{location}" site:brightnetwork.co.uk',
            f'"{role}" "{location}" site:indeed.com',
            f'"{role}" "{location}" site:linkedin.com/jobs',
            f'"{role}" "{location}" careers apply',
        ])

    if extra_keywords:
        queries.append(f'{extra_keywords} "{location}" graduate jobs careers apply')

    # Remove duplicates while preserving order.
    seen = set()
    unique = []
    for q in queries:
        if q not in seen:
            seen.add(q)
            unique.append(q)

    return unique


def infer_company_from_title(title: str) -> str:
    """
    Very rough company inference from search result title.
    Better versions should use an LLM extractor.
    """
    if not title:
        return ""

    separators = [" - ", " | ", " at ", " – ", " — "]

    for sep in separators:
        if sep in title:
            parts = title.split(sep)
            # Usually company is last part on job boards, but not always.
            candidate = parts[-1].strip()
            if len(candidate.split()) <= 5:
                return candidate

    return ""


def extract_job_from_search_result(result: dict) -> dict:
    title = result.get("title", "")
    snippet = result.get("snippet", "")

    return {
        "job_title": title,
        "company": infer_company_from_title(title),
        "location": "",
        "source_url": result.get("url", ""),
        "source_domain": result.get("source", ""),
        "snippet": snippet,
        "description": snippet,
        "query": result.get("query", ""),
    }


def deduplicate_jobs(jobs: list) -> list:
    unique = []
    seen_urls = set()

    for job in jobs:
        url = job.get("source_url", "")
        if not url or url in seen_urls:
            continue

        seen_urls.add(url)
        unique.append(job)

    return unique


# ============================================================
# JOB PAGE FETCHING
# ============================================================

def fetch_page_text(url: str) -> str:
    headers = {
        "User-Agent": "Mozilla/5.0 JobFinderAgent/1.0"
    }

    try:
        response = requests.get(url, headers=headers, timeout=15)
        if response.status_code >= 400:
            return ""
        soup = BeautifulSoup(response.text, "html.parser")

        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        text = soup.get_text(" ")
        return clean_text(text)
    except Exception:
        return ""


def enrich_job_description(job: dict) -> dict:
    page_text = fetch_page_text(job.get("source_url", ""))

    if page_text and len(page_text) > len(job.get("description", "")):
        job["description"] = truncate(page_text, 8000)

    return job


# ============================================================
# COMPANY CAREER VERIFICATION
# ============================================================

CAREER_KEYWORDS = [
    "careers",
    "jobs",
    "workday",
    "greenhouse",
    "lever",
    "smartrecruiters",
    "myworkdayjobs",
    "ashbyhq",
    "workable",
    "icims",
    "successfactors",
]


def looks_like_career_url(url: str) -> bool:
    lower = url.lower()
    return any(k in lower for k in CAREER_KEYWORDS)


def verify_job_on_company_site(job: dict) -> dict:
    title = job.get("job_title", "")
    company = job.get("company", "")
    source_domain = job.get("source_domain", "")

    query_parts = []

    if company:
        query_parts.append(f'"{company}"')

    if title:
        simplified_title = title[:100]
        query_parts.append(f'"{simplified_title}"')

    query_parts.append("careers OR jobs OR workday OR greenhouse OR lever apply")

    query = " ".join(query_parts)

    results = serpapi_search(query, num_results=8)

    best = None
    best_score = 0

    for r in results:
        url = r.get("url", "")
        r_title = r.get("title", "")
        snippet = r.get("snippet", "")

        score = 0

        title_similarity = fuzz.partial_ratio(title.lower(), (r_title + " " + snippet).lower())

        if title_similarity > 80:
            score += 40
        elif title_similarity > 60:
            score += 25
        elif title_similarity > 40:
            score += 10

        if company and company.lower() in (r_title + " " + snippet + " " + url).lower():
            score += 25

        if looks_like_career_url(url):
            score += 25

        if source_domain and source_domain in get_domain(url):
            score += 10

        if score > best_score:
            best_score = score
            best = r

    if best_score >= 75:
        status = "Verified"
    elif best_score >= 55:
        status = "Likely real"
    elif best_score >= 35:
        status = "Unverified"
    else:
        status = "Low confidence"

    return {
        "verification_status": status,
        "verification_score": best_score,
        "official_url": best.get("url", "") if best else "",
        "verification_evidence": best.get("title", "") if best else "",
    }


# ============================================================
# MATCH SCORING
# ============================================================

def keyword_match_score(job_text: str, cv_text: str) -> int:
    job_words = set(re.findall(r"[a-zA-Z][a-zA-Z\+\#\.\-]{2,}", job_text.lower()))
    cv_words = set(re.findall(r"[a-zA-Z][a-zA-Z\+\#\.\-]{2,}", cv_text.lower()))

    stopwords = {
        "the", "and", "for", "with", "you", "are", "this", "that", "will",
        "from", "have", "has", "your", "our", "their", "job", "role",
        "candidate", "company", "work", "team", "skills"
    }

    job_words = {w for w in job_words if w not in stopwords and len(w) > 2}

    if not job_words:
        return 0

    overlap = job_words.intersection(cv_words)
    score = int((len(overlap) / max(len(job_words), 1)) * 100)

    return min(score, 100)


def llm_match_analysis(job: dict, profile: dict, cv_text: str) -> dict:
    prompt = f"""
You are a career matching assistant.

Score how well this candidate fits this job.

Return ONLY valid JSON:

{{
  "match_score": 0,
  "fit_summary": "",
  "matched_skills": [],
  "missing_skills": [],
  "best_projects_to_highlight": [],
  "cv_tailoring_strategy": []
}}

Rules:
- Score from 0 to 100.
- Do not invent candidate skills.
- Be honest.
- Focus on UK graduate/intern/entry-level roles.

Candidate profile:
{json.dumps(profile, indent=2)}

Candidate CV text:
\"\"\"
{truncate(cv_text, 6000)}
\"\"\"

Job:
{json.dumps(job, indent=2)}
"""

    result = call_openai(prompt)
    parsed = extract_json_from_text(result)

    if parsed:
        return parsed

    fallback_score = keyword_match_score(
        job.get("description", "") + " " + job.get("job_title", ""),
        cv_text
    )

    return {
        "match_score": fallback_score,
        "fit_summary": "Keyword-based fallback score. Add OpenAI API key for deeper analysis.",
        "matched_skills": [],
        "missing_skills": [],
        "best_projects_to_highlight": [],
        "cv_tailoring_strategy": [],
    }


# ============================================================
# DOCUMENT GENERATION
# ============================================================

def generate_tailored_cv_text(job: dict, profile: dict, cv_text: str, match: dict) -> str:
    prompt = f"""
You are an expert UK CV writer.

Create a tailored ATS-friendly CV for this specific job.

Important rules:
- Do NOT invent experience, education, skills, dates, employers, projects, or results.
- Only use information from the candidate CV/profile.
- You may rewrite and reorder bullets to match the job.
- Keep it professional and honest.
- Use British English.
- Make it suitable for graduate/early-career AI, data science, ML, or technology roles.
- Keep it concise enough for 1-2 pages.

Return the CV in clean plain text with these sections:
1. Name and Contact
2. Professional Profile
3. Education
4. Technical Skills
5. Projects
6. Experience
7. Certifications / Additional Information

Candidate profile:
{json.dumps(profile, indent=2)}

Original CV:
\"\"\"
{truncate(cv_text, 10000)}
\"\"\"

Job:
{json.dumps(job, indent=2)}

Match analysis:
{json.dumps(match, indent=2)}
"""

    result = call_openai(prompt)

    if result and not result.startswith("OPENAI_ERROR"):
        return result

    return f"""
TAILORED CV DRAFT

Professional Profile
AI and data science candidate with experience in Python, machine learning, NLP, data analysis, and cybersecurity. This CV should be manually refined because the OpenAI API was unavailable.

Relevant Skills
{", ".join(profile.get("technical_skills", []))}

Relevant Projects
{json.dumps(profile.get("projects", []), indent=2)}

Experience
{json.dumps(profile.get("experience", []), indent=2)}

Education
{json.dumps(profile.get("education", []), indent=2)}
"""


def generate_cover_letter_text(job: dict, profile: dict, cv_text: str, match: dict) -> str:
    prompt = f"""
Write a concise UK-style cover letter for this job.

Rules:
- Do not invent anything.
- Use British English.
- Keep it around 250-350 words.
- Make it specific to the role.
- Mention relevant projects and skills honestly.
- Do not sound exaggerated.

Candidate profile:
{json.dumps(profile, indent=2)}

Original CV:
\"\"\"
{truncate(cv_text, 8000)}
\"\"\"

Job:
{json.dumps(job, indent=2)}

Match analysis:
{json.dumps(match, indent=2)}
"""

    result = call_openai(prompt)

    if result and not result.startswith("OPENAI_ERROR"):
        return result

    return f"""
Dear Hiring Manager,

I am writing to express my interest in the {job.get("job_title", "role")} position. My background combines artificial intelligence, machine learning, data analysis, and cybersecurity, with hands-on academic and project experience in Python and applied AI.

I am particularly interested in this opportunity because it aligns with my technical background and my goal of developing practical AI and data-driven solutions. My projects and academic work have involved machine learning, NLP, data processing, and model evaluation, which I believe are relevant to this role.

I would welcome the opportunity to contribute my skills and continue developing within your team.

Kind regards,
{profile.get("name", "")}
"""


def save_text_as_docx(title: str, text: str, filename: str) -> str:
    doc = Document()

    doc.add_heading(title, level=1)

    for block in text.split("\n"):
        line = block.strip()
        if not line:
            continue

        if len(line) < 70 and not line.endswith(".") and not line.endswith(","):
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)

    return path


def save_application_pack(job: dict, match: dict, cv_text: str, cover_letter: str) -> dict:
    company = job.get("company") or "Company"
    role = job.get("job_title") or "Role"

    base = safe_filename(f"{company}_{role}_{datetime.now().strftime('%Y%m%d_%H%M%S')}")

    cv_path = save_text_as_docx(
        title=f"Tailored CV - {role}",
        text=cv_text,
        filename=f"{base}_CV.docx"
    )

    cover_path = save_text_as_docx(
        title=f"Cover Letter - {role}",
        text=cover_letter,
        filename=f"{base}_Cover_Letter.docx"
    )

    summary_doc = Document()
    summary_doc.add_heading("Application Pack Summary", level=1)

    summary_doc.add_heading("Job Details", level=2)
    summary_doc.add_paragraph(f"Role: {role}")
    summary_doc.add_paragraph(f"Company: {company}")
    summary_doc.add_paragraph(f"Source URL: {job.get('source_url', '')}")
    summary_doc.add_paragraph(f"Official/Verification URL: {job.get('official_url', '')}")
    summary_doc.add_paragraph(f"Verification Status: {job.get('verification_status', '')}")
    summary_doc.add_paragraph(f"Verification Score: {job.get('verification_score', '')}")
    summary_doc.add_paragraph(f"Match Score: {match.get('match_score', '')}")

    summary_doc.add_heading("Fit Summary", level=2)
    summary_doc.add_paragraph(match.get("fit_summary", ""))

    summary_doc.add_heading("Matched Skills", level=2)
    for skill in match.get("matched_skills", []):
        summary_doc.add_paragraph(str(skill), style="List Bullet")

    summary_doc.add_heading("Missing Skills", level=2)
    for skill in match.get("missing_skills", []):
        summary_doc.add_paragraph(str(skill), style="List Bullet")

    summary_path = os.path.join(OUTPUT_DIR, f"{base}_Summary.docx")
    summary_doc.save(summary_path)

    return {
        "cv_path": cv_path,
        "cover_letter_path": cover_path,
        "summary_path": summary_path,
    }


# ============================================================
# STREAMLIT UI
# ============================================================

st.set_page_config(
    page_title="Job Finder AI Agent",
    page_icon="💼",
    layout="wide"
)

st.title("💼 Job Finder AI Agent")
st.caption("Find real jobs, verify them on company career sites, score fit, and generate tailored application documents.")

with st.sidebar:
    st.header("Settings")

    st.write("API status:")

    if OPENAI_API_KEY:
        st.success("OpenAI API key loaded")
    else:
        st.warning("OpenAI API key missing")

    if SERPAPI_API_KEY:
        st.success("SerpAPI key loaded")
    else:
        st.warning("SerpAPI key missing")

    st.divider()

    max_queries = st.slider("Max search queries", 1, 20, 5)
    results_per_query = st.slider("Results per query", 3, 20, 8)
    max_jobs_to_process = st.slider("Max jobs to process", 3, 30, 10)

    st.divider()

    st.info(
        "This MVP does not log into LinkedIn or auto-apply. "
        "It finds and verifies jobs, then gives you the documents to apply manually."
    )


uploaded_cv = st.file_uploader("Upload your master CV", type=["pdf", "docx", "txt"])

col1, col2 = st.columns(2)

with col1:
    target_roles = st.text_area(
        "Target roles",
        value="Data Scientist Graduate\nMachine Learning Engineer Graduate\nAI Engineer Graduate\nNLP Engineer Intern\nCybersecurity Data Analyst",
        height=160
    )

with col2:
    location = st.text_input("Location", value="London UK")
    extra_keywords = st.text_input(
        "Extra keywords",
        value="Python machine learning NLP graduate"
    )

run_button = st.button("🚀 Run Job Finder Agent", type="primary")


if run_button:
    if not uploaded_cv:
        st.error("Please upload your CV first.")
        st.stop()

    if not SERPAPI_API_KEY:
        st.error("Please add SERPAPI_API_KEY to your .env file.")
        st.stop()

    # Step 1: Read CV
    with st.spinner("Reading your CV..."):
        cv_text = read_uploaded_cv(uploaded_cv)

    st.success("CV loaded successfully.")

    with st.expander("Preview extracted CV text"):
        st.write(cv_text[:4000])

    # Step 2: Build profile
    with st.spinner("Building candidate profile..."):
        profile = build_candidate_profile(cv_text)

    st.subheader("Candidate Profile")
    st.json(profile)

    # Step 3: Generate job search queries
    queries = generate_job_queries(target_roles, location, extra_keywords)
    queries = queries[:max_queries]

    st.subheader("Search Queries")
    for q in queries:
        st.code(q)

    # Step 4: Search jobs
    all_results = []

    with st.spinner("Searching job sources..."):
        for q in queries:
            results = serpapi_search(q, num_results=results_per_query)
            all_results.extend(results)
            time.sleep(0.5)

    if not all_results:
        st.error("No search results found.")
        st.stop()

    raw_jobs = [extract_job_from_search_result(r) for r in all_results]
    jobs = deduplicate_jobs(raw_jobs)
    jobs = jobs[:max_jobs_to_process]

    st.success(f"Found {len(jobs)} unique candidate jobs.")

    # Step 5: Process jobs
    processed_jobs = []

    progress = st.progress(0)

    for idx, job in enumerate(jobs):
        with st.spinner(f"Processing job {idx + 1}/{len(jobs)}: {job.get('job_title', '')[:80]}"):
            job = enrich_job_description(job)

            verification = verify_job_on_company_site(job)
            job.update(verification)

            match = llm_match_analysis(job, profile, cv_text)

            job["match_score"] = match.get("match_score", 0)
            job["fit_summary"] = match.get("fit_summary", "")
            job["matched_skills"] = ", ".join(match.get("matched_skills", []))
            job["missing_skills"] = ", ".join(match.get("missing_skills", []))
            job["_match_json"] = match

            processed_jobs.append(job)

        progress.progress((idx + 1) / len(jobs))

    # Step 6: Show results
    df = pd.DataFrame([
        {
            "Job Title": j.get("job_title", ""),
            "Company": j.get("company", ""),
            "Match Score": j.get("match_score", 0),
            "Verification": j.get("verification_status", ""),
            "Verification Score": j.get("verification_score", 0),
            "Source": j.get("source_domain", ""),
            "Source URL": j.get("source_url", ""),
            "Official URL": j.get("official_url", ""),
            "Fit Summary": j.get("fit_summary", ""),
            "Matched Skills": j.get("matched_skills", ""),
            "Missing Skills": j.get("missing_skills", ""),
        }
        for j in processed_jobs
    ])

    df = df.sort_values(by=["Verification Score", "Match Score"], ascending=False)

    st.subheader("Ranked Jobs")
    st.dataframe(df, use_container_width=True)

    csv = df.to_csv(index=False).encode("utf-8")
    st.download_button(
        label="Download job results as CSV",
        data=csv,
        file_name="job_results.csv",
        mime="text/csv"
    )

    # Step 7: Select job for document generation
    st.subheader("Generate Application Pack")

    options = []
    for i, j in enumerate(processed_jobs):
        label = f"{i}: {j.get('job_title', '')[:80]} | Match {j.get('match_score', 0)} | {j.get('verification_status', '')}"
        options.append(label)

    selected = st.selectbox("Choose a job", options)

    selected_index = int(selected.split(":")[0])
    selected_job = processed_jobs[selected_index]
    selected_match = selected_job["_match_json"]

    st.write("Selected job:")
    st.json({
        "job_title": selected_job.get("job_title", ""),
        "company": selected_job.get("company", ""),
        "source_url": selected_job.get("source_url", ""),
        "official_url": selected_job.get("official_url", ""),
        "verification_status": selected_job.get("verification_status", ""),
        "match_score": selected_match.get("match_score", 0),
        "fit_summary": selected_match.get("fit_summary", ""),
    })

    if st.button("📄 Generate tailored CV + cover letter"):
        with st.spinner("Generating tailored CV..."):
            tailored_cv = generate_tailored_cv_text(
                selected_job,
                profile,
                cv_text,
                selected_match
            )

        with st.spinner("Generating cover letter..."):
            cover_letter = generate_cover_letter_text(
                selected_job,
                profile,
                cv_text,
                selected_match
            )

        with st.spinner("Saving documents..."):
            paths = save_application_pack(
                selected_job,
                selected_match,
                tailored_cv,
                cover_letter
            )

        st.success("Application pack generated.")

        st.subheader("Tailored CV Preview")
        st.text_area("CV", tailored_cv, height=400)

        st.subheader("Cover Letter Preview")
        st.text_area("Cover Letter", cover_letter, height=300)

        with open(paths["cv_path"], "rb") as f:
            st.download_button(
                "Download Tailored CV",
                data=f,
                file_name=os.path.basename(paths["cv_path"]),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with open(paths["cover_letter_path"], "rb") as f:
            st.download_button(
                "Download Cover Letter",
                data=f,
                file_name=os.path.basename(paths["cover_letter_path"]),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with open(paths["summary_path"], "rb") as f:
            st.download_button(
                "Download Application Summary",
                data=f,
                file_name=os.path.basename(paths["summary_path"]),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )